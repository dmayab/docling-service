#!/usr/bin/env python3
"""
Script para crear un documento DOCX de prueba para verificar el soporte de Docling
"""

from docx import Document
import os

def create_test_docx():
    """Crea un documento DOCX de prueba"""
    
    print("📝 Creando documento DOCX de prueba...")
    
    # Crear documento
    doc = Document()
    
    # Título
    title = doc.add_heading('Documento de Prueba DOCX', 0)
    
    # Párrafo de introducción
    intro = doc.add_paragraph('Este es un documento de prueba para verificar que Docling puede procesar archivos DOCX correctamente.')
    
    # Lista
    doc.add_heading('Lista de Verificación:', level=1)
    doc.add_paragraph('• Extracción de texto funcional', style='List Bullet')
    doc.add_paragraph('• Estructura del documento preservada', style='List Bullet')
    doc.add_paragraph('• Elementos sin coordenadas físicas', style='List Bullet')
    doc.add_paragraph('• Búsqueda de texto habilitada', style='List Bullet')
    
    # Sección con más contenido
    doc.add_heading('Contenido de Ejemplo', level=1)
    
    content_para = doc.add_paragraph()
    content_para.add_run('Este párrafo contiene ').bold = False
    content_para.add_run('texto en negrita').bold = True
    content_para.add_run(' y también ').bold = False
    content_para.add_run('texto en cursiva').italic = True
    content_para.add_run(' para probar el procesamiento de diferentes estilos.').bold = False
    
    # Otra sección
    doc.add_heading('Información Técnica', level=2)
    doc.add_paragraph('Los archivos DOCX no mantienen coordenadas físicas como los PDF, pero Docling puede extraer la estructura del documento y el contenido textual completo.')
    
    # Párrafo final
    doc.add_paragraph('Este documento debería procesarse correctamente y permitir búsquedas de texto, aunque sin highlighting visual debido a la ausencia de coordenadas físicas.')
    
    # Guardar archivo
    filename = 'test_document.docx'
    filepath = os.path.join(os.path.dirname(__file__), filename)
    doc.save(filepath)
    
    print(f"✅ Documento DOCX creado: {filepath}")
    print(f"📄 Archivo listo para prueba con Docling")
    
    return filepath

if __name__ == "__main__":
    create_test_docx()